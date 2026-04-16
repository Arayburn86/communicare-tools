import streamlit as st
import anthropic
import os
import datetime
import io

st.set_page_config(page_title="Grant Narrative Generator", page_icon="✍️", layout="wide")

st.markdown("""
<style>
h1{color:#1A6E6E!important}h2{color:#1A6E6E!important;font-size:1.1rem!important}
h3{color:#1F2937!important;font-size:1rem!important}
.stButton>button{background:#1A6E6E!important;color:white!important;border:none!important;
  border-radius:8px!important;font-weight:600!important}
.hint{background:#E1F5EE;border-radius:8px;padding:10px 14px;
  font-size:.85rem;color:#085041;margin-bottom:10px}
.warn{background:#FEF3C7;border-left:3px solid #BA7517;border-radius:0 8px 8px 0;
  padding:10px 14px;font-size:.85rem;color:#92400E;margin-bottom:10px}
.section-head{font-weight:600;color:#1A6E6E;font-size:1rem;
  margin:20px 0 6px;padding-bottom:4px;border-bottom:2px solid #E1F5EE}
.word-count{font-size:.75rem;color:#9CA3AF;margin-bottom:4px}
</style>""", unsafe_allow_html=True)


def get_api_key():
    if "ANTHROPIC_API_KEY" in st.secrets:
        return st.secrets["ANTHROPIC_API_KEY"]
    return os.environ.get("ANTHROPIC_API_KEY", "")


# ─────────────────────────────────────────────────────────────────────────────
# Section generation
# ─────────────────────────────────────────────────────────────────────────────

SYSTEM = """You are an expert grant writer for a nonprofit organization.
You write clear, professional, compelling grant narratives for foundation funders.
Your tone is warm but precise — you lead with community need, ground everything
in data, and connect program activities directly to outcomes.
Write in connected prose paragraphs. No bullet points. No markdown. No section headings."""


def generate_section(client, prompt):
    msg = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=900,
        system=SYSTEM,
        messages=[{"role": "user", "content": prompt}],
    )
    return msg.content[0].text.strip()


def build_prompts(p):
    """Build all 6 section prompts from program data dict."""
    base = (
        f"You are drafting a grant narrative for {p['org_name']}, a nonprofit in "
        f"{p['org_location']} founded in {p['org_founded']}. "
        f"Mission: {p['org_mission']}. "
        f"Program: {p['program_name']}. "
        f"Funder: {p['funder_name']} (priority areas: {p['funder_priority']}). "
        f"Requesting {p['grant_amount']} for {p['grant_period']}."
    )

    stats = "\n".join(f"- {s}" for s in p["community_stats"])
    acts  = "\n".join(f"- {a}" for a in p["key_activities"])
    short = "\n".join(f"- {o}" for o in p["short_term_outcomes"])
    long_ = "\n".join(f"- {o}" for o in p["long_term_outcomes"])
    staff = "\n".join(f"- {s}" for s in p["key_staff"])
    parts = "\n".join(f"- {pt}" for pt in p["partnerships"])

    return {
        "Executive Summary": {
            "target": "150–200",
            "prompt": (
                f"{base}\n\nWrite a concise executive summary (150–200 words). "
                "Introduce the organization, name the community need, describe the program, "
                "state the funding request, and summarize expected impact. "
                "Start with the organization name."
            ),
        },
        "Statement of Need": {
            "target": "300–400",
            "prompt": (
                f"{base}\n\nWrite a compelling Statement of Need (300–400 words). "
                f"Target population: {p['target_population']}.\n\n"
                f"Use these data points:\n{stats}\n\n"
                "Connect them into a narrative about why this program is urgently needed. "
                "Emphasize health equity and the gap between need and available services."
            ),
        },
        "Program Description": {
            "target": "350–450",
            "prompt": (
                f"{base}\n\nWrite a Program Description (350–450 words).\n\n"
                f"Program overview: {p['program_description']}\n"
                f"Evidence base: {p['evidence_base']}\n"
                f"Key activities:\n{acts}\n\n"
                "Describe how the program works, why these activities were chosen, "
                "and how the evidence base supports this approach. "
                "Emphasize cultural responsiveness and community-centered design."
            ),
        },
        "Goals, Objectives & Evaluation": {
            "target": "300–400",
            "prompt": (
                f"{base}\n\nWrite a Goals, Objectives & Evaluation section (300–400 words).\n\n"
                f"Short-term outcomes:\n{short}\n\n"
                f"Long-term outcomes:\n{long_}\n\n"
                f"Evaluation approach: {p['evaluation_methods']}\n\n"
                "Frame outcomes as SMART objectives where possible. "
                "Describe the evaluation plan and how data will drive program improvement."
            ),
        },
        "Organizational Capacity": {
            "target": "250–350",
            "prompt": (
                f"{base}\n\nWrite an Organizational Capacity section (250–350 words).\n\n"
                f"Key staff:\n{staff}\n\n"
                f"Community partnerships:\n{parts}\n\n"
                f"Past success: {p['past_success']}\n\n"
                "Demonstrate that the organization has the staff expertise, community "
                "relationships, and track record to successfully implement this program. "
                "Reference pilot data and partnerships specifically."
            ),
        },
        "Sustainability Plan": {
            "target": "200–280",
            "prompt": (
                f"{base}\n\nWrite a Sustainability Plan (200–280 words). "
                "Address how the program will continue after this grant period ends. "
                "Include: Medicaid billing, other foundation grants, state contracts, "
                "individual donors, plans to build the evidence base for continued "
                "investment, and community ownership. Be realistic and specific."
            ),
        },
    }


# ─────────────────────────────────────────────────────────────────────────────
# UI
# ─────────────────────────────────────────────────────────────────────────────

st.title("✍️ Grant Narrative Generator")
st.caption("Fill in your program details — Claude drafts all 6 sections of your grant narrative.")

api_key = get_api_key()
if not api_key:
    st.markdown(
        '<div class="warn">⚠ API key not set. Add ANTHROPIC_API_KEY in '
        'Streamlit Cloud → Settings → Secrets.</div>',
        unsafe_allow_html=True,
    )

if "narrative_sections" not in st.session_state:
    st.session_state.narrative_sections = {}
if "narrative_generated" not in st.session_state:
    st.session_state.narrative_generated = False

# ── Input form ────────────────────────────────────────────────────────────────

with st.expander("## Program & funder details", expanded=True):
    c1, c2 = st.columns(2)
    with c1:
        org_name      = st.text_input("Organization name", "Communicare Alliance")
        org_location  = st.text_input("Location",          "Woonsocket, RI")
        org_founded   = st.text_input("Year founded",      "1994")
        org_mission   = st.text_area("Mission statement",
            "To strengthen the health and well-being of individuals and families in "
            "Woonsocket and the surrounding Blackstone Valley through community-centered, "
            "culturally responsive human services.", height=80)
    with c2:
        program_name  = st.text_input("Program name",    "Youth Mental Health & Resilience Program")
        funder_name   = st.text_input("Funder name",     "Rhode Island Foundation")
        funder_priority = st.text_input("Funder priority areas",
            "youth mental health, health equity, community resilience")
        grant_amount  = st.text_input("Amount requested", "$185,000")
        grant_period  = st.text_input("Grant period",    "January 2025 – December 2026")

with st.expander("## Target population & community need", expanded=True):
    target_pop = st.text_input("Target population",
        "Youth ages 12–18 in Woonsocket, RI")

    st.markdown("**Community need data points** — one per line")
    community_stats = st.text_area("Community stats",
        "Woonsocket has a youth poverty rate of 38%, nearly triple the state average\n"
        "53% of youth qualify for free or reduced-price lunch\n"
        "Rates of adolescent depression are 22% above the national average\n"
        "Only 1 in 5 youth who need mental health services currently receives them\n"
        "The population is 42% Latino and 12% Black — communities historically underserved by behavioral health systems",
        height=120)

with st.expander("## Program design", expanded=True):
    program_desc = st.text_area("Program description (2-3 sentences)",
        "A two-year school- and community-based program delivering individual counseling, "
        "resilience skills groups, family psychoeducation, peer mentorship, and trauma-informed "
        "staff training to 150 youth annually.",
        height=80)

    evidence_base = st.text_area("Evidence base",
        "The program uses Cognitive Behavioral Therapy for Adolescents (CBT-A) and the "
        "COPE curriculum, both with strong evidence bases in peer-reviewed literature for "
        "reducing adolescent anxiety and depression.",
        height=70)

    st.markdown("**Key activities** — one per line")
    key_activities = st.text_area("Activities",
        "Weekly individual counseling sessions (school-based and clinic)\n"
        "Bi-weekly 8-week resilience skills groups (COPE curriculum, 10 youth per cohort)\n"
        "Monthly family psychoeducation workshops in English and Spanish\n"
        "Annual school staff training on trauma-informed practices\n"
        "Peer mentor recruitment and training (15 mentors per year)",
        height=110)

with st.expander("## Outcomes & evaluation", expanded=False):
    st.markdown("**Short-term outcomes** — one per line")
    short_outcomes = st.text_area("Short-term outcomes",
        "75% of discharged youth show clinically meaningful PHQ-A reduction (≥5 points)\n"
        "80% of group completers show improved coping skills on COPE outcomes scale\n"
        "70% of caregiver workshop attendees report increased confidence",
        height=90)

    st.markdown("**Long-term outcomes** — one per line")
    long_outcomes = st.text_area("Long-term outcomes",
        "Sustained reduction in adolescent mental health crisis events in Woonsocket\n"
        "Improved school attendance and academic engagement among participants\n"
        "A replicable, evidence-informed community mental health model for regional scale",
        height=80)

    eval_methods = st.text_area("Evaluation approach",
        "Pre/post validated screening tools (PHQ-A, GAD-7, COPE outcomes scale), "
        "school attendance data from partner districts, quarterly data review with staff, "
        "and an independent Year 2 evaluation by URI School of Public Health.",
        height=70)

with st.expander("## Organizational capacity", expanded=False):
    st.markdown("**Key staff** — one per line")
    key_staff = st.text_area("Staff",
        "Director of Programs (MSW, LCSW, 12 years community behavioral health)\n"
        "Two Licensed Clinical Social Workers (LCSW)\n"
        "Community Health Educator (bilingual Spanish/English)\n"
        "Program Coordinator",
        height=90)

    st.markdown("**Key partnerships** — one per line")
    partnerships = st.text_area("Partnerships",
        "Woonsocket Education Department (3 partner schools)\n"
        "Thundermist Health Center\n"
        "RI Department of Children, Youth & Families (DCYF)\n"
        "Our Lady of Fatima Parish Community",
        height=90)

    past_success = st.text_area("Past success / pilot data",
        "In our 2022–2024 pilot (n=60 youth), 78% showed clinically meaningful PHQ-A "
        "reduction after 8 weeks; school attendance improved by an average of 11 days "
        "per participant among youth with chronic absenteeism.",
        height=70)

# ── Generate button ───────────────────────────────────────────────────────────

st.markdown("---")
g1, g2 = st.columns([2, 3])
with g1:
    generate_btn = st.button(
        "✨ Generate all 6 sections",
        disabled=not api_key,
        use_container_width=True,
    )
with g2:
    st.markdown(
        '<div class="hint" style="margin:0">Takes about 60–90 seconds · '
        'costs ~$0.05 in API fees · '
        'review every section before submitting</div>',
        unsafe_allow_html=True,
    )

if generate_btn:
    program = {
        "org_name":          org_name,
        "org_location":      org_location,
        "org_founded":       org_founded,
        "org_mission":       org_mission,
        "program_name":      program_name,
        "funder_name":       funder_name,
        "funder_priority":   funder_priority,
        "grant_amount":      grant_amount,
        "grant_period":      grant_period,
        "target_population": target_pop,
        "community_stats":   [s.strip() for s in community_stats.split("\n") if s.strip()],
        "program_description": program_desc,
        "evidence_base":     evidence_base,
        "key_activities":    [a.strip() for a in key_activities.split("\n") if a.strip()],
        "short_term_outcomes": [o.strip() for o in short_outcomes.split("\n") if o.strip()],
        "long_term_outcomes":  [o.strip() for o in long_outcomes.split("\n") if o.strip()],
        "evaluation_methods":  eval_methods,
        "key_staff":         [s.strip() for s in key_staff.split("\n") if s.strip()],
        "partnerships":      [p.strip() for p in partnerships.split("\n") if p.strip()],
        "past_success":      past_success,
    }

    prompts = build_prompts(program)
    client  = anthropic.Anthropic(api_key=api_key)
    sections = {}
    progress = st.progress(0, text="Drafting executive summary...")

    for i, (section_name, meta) in enumerate(prompts.items()):
        progress.progress(
            (i + 1) / len(prompts),
            text=f"Drafting {section_name}...",
        )
        sections[section_name] = {
            "text":   generate_section(client, meta["prompt"]),
            "target": meta["target"],
        }

    progress.empty()
    st.session_state.narrative_sections  = sections
    st.session_state.narrative_generated = True
    st.session_state.narrative_program   = program
    st.rerun()

# ── Display & edit generated narrative ───────────────────────────────────────

if st.session_state.narrative_generated and st.session_state.narrative_sections:
    sec = st.session_state.narrative_sections
    p   = st.session_state.get("narrative_program", {})
    today = datetime.date.today().strftime("%B %d, %Y")

    st.markdown("---")
    st.markdown("## Your Grant Narrative")
    st.markdown(
        f"**{p.get('program_name','')}** — {p.get('funder_name','')} "
        f"· {p.get('grant_amount','')} · {today}"
    )
    st.markdown(
        '<div class="warn">📝 AI-assisted draft — review and edit every section '
        'before submitting to a funder. Adjust tone, add specifics, '
        'verify all data points.</div>',
        unsafe_allow_html=True,
    )

    # Editable sections
    edited = {}
    for section_name, content in sec.items():
        text   = content["text"]
        target = content["target"]
        wc     = len(text.split())

        st.markdown(
            f'<div class="section-head">{section_name}</div>',
            unsafe_allow_html=True,
        )
        st.markdown(
            f'<div class="word-count">Target: {target} words · '
            f'Current: {wc} words</div>',
            unsafe_allow_html=True,
        )

        edited_text = st.text_area(
            section_name,
            value=text,
            height=200,
            key=f"edit_{section_name}",
            label_visibility="collapsed",
        )
        edited[section_name] = edited_text

        # Regenerate individual section
        if st.button(f"🔄 Regenerate this section",
                     key=f"regen_{section_name}"):
            prompts = build_prompts(p)
            client  = anthropic.Anthropic(api_key=api_key)
            with st.spinner(f"Rewriting {section_name}..."):
                new_text = generate_section(
                    client, prompts[section_name]["prompt"]
                )
            st.session_state.narrative_sections[section_name]["text"] = new_text
            st.rerun()

    # ── Download as Word doc ──────────────────────────────────────────────────

    st.markdown("---")
    st.markdown("### Download")

    # Build plain-text version first (always available)
    full_text = f"""{p.get('program_name', 'Grant Narrative')}
Grant Narrative — {p.get('funder_name', '')}
{p.get('org_name', '')} · {p.get('grant_amount', '')} · {p.get('grant_period', '')}
Prepared: {today}

{'='*60}
"""
    for section_name, text in edited.items():
        full_text += f"\n{section_name.upper()}\n{'─'*40}\n{text}\n"

    full_text += f"\n{'='*60}\nAI-assisted draft — review before submitting to funder\n"

    c1, c2 = st.columns(2)

    with c1:
        st.download_button(
            "⬇ Download as text file",
            data=full_text.encode("utf-8"),
            file_name=f"grant_narrative_{p.get('funder_name','').replace(' ','_')}_{today}.txt",
            mime="text/plain",
            use_container_width=True,
        )

    with c2:
        # Word doc via python-docx
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            sec_props = doc.sections[0]
            sec_props.page_width  = Inches(8.5)
            sec_props.page_height = Inches(11)
            sec_props.left_margin = sec_props.right_margin = \
                sec_props.top_margin = sec_props.bottom_margin = Inches(1)

            TEAL = RGBColor(0x1A, 0x6E, 0x6E)
            DARK = RGBColor(0x1F, 0x29, 0x37)
            GRAY = RGBColor(0xAA, 0xAA, 0xAA)
            RED  = RGBColor(0xAA, 0x44, 0x44)

            # Title
            tp = doc.add_paragraph()
            tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tr = tp.add_run(p.get("program_name", "Grant Narrative"))
            tr.bold = True; tr.font.size = Pt(18)
            tr.font.color.rgb = TEAL; tr.font.name = "Arial"

            sp = doc.add_paragraph()
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sp.paragraph_format.space_before = Pt(2)
            sp.paragraph_format.space_after  = Pt(2)
            sr = sp.add_run(
                f"{p.get('org_name','')}  |  Submitted to: {p.get('funder_name','')}  "
                f"|  {p.get('grant_amount','')}  |  {p.get('grant_period','')}"
            )
            sr.font.size = Pt(9); sr.font.color.rgb = GRAY; sr.font.name = "Arial"

            dp = doc.add_paragraph()
            dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            dp.paragraph_format.space_after = Pt(14)
            dr = dp.add_run(f"Generated: {today}  |  AI-Assisted Draft — Review Before Submission")
            dr.italic = True; dr.font.size = Pt(9)
            dr.font.color.rgb = RED; dr.font.name = "Arial"

            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            def add_divider():
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after  = Pt(2)
                pPr  = para._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bot  = OxmlElement("w:bottom")
                bot.set(qn("w:val"),   "single")
                bot.set(qn("w:sz"),    "4")
                bot.set(qn("w:space"), "1")
                bot.set(qn("w:color"), "9FE1CB")
                pBdr.append(bot)
                pPr.append(pBdr)

            add_divider()

            for section_name, text in edited.items():
                # Heading
                hp = doc.add_paragraph()
                hp.paragraph_format.space_before = Pt(14)
                hp.paragraph_format.space_after  = Pt(4)
                hr = hp.add_run(section_name.upper())
                hr.bold = True; hr.font.size = Pt(12)
                hr.font.color.rgb = TEAL; hr.font.name = "Arial"

                # Target word count
                wcp = doc.add_paragraph()
                wcp.paragraph_format.space_before = Pt(0)
                wcp.paragraph_format.space_after  = Pt(6)
                wcr = wcp.add_run(
                    f"Target: {sec[section_name]['target']} words  ·  "
                    f"Current: {len(text.split())} words"
                )
                wcr.italic = True; wcr.font.size = Pt(8)
                wcr.font.color.rgb = GRAY; wcr.font.name = "Arial"

                # Body paragraphs
                for block in text.strip().split("\n\n"):
                    block = block.strip()
                    if block:
                        bp = doc.add_paragraph()
                        bp.paragraph_format.space_before = Pt(0)
                        bp.paragraph_format.space_after  = Pt(8)
                        br = bp.add_run(block.replace("\n", " "))
                        br.font.size = Pt(11)
                        br.font.color.rgb = DARK; br.font.name = "Arial"

                add_divider()

            # Footer
            fp = doc.add_paragraph()
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fp.paragraph_format.space_before = Pt(12)
            fr = fp.add_run(
                f"{p.get('org_name','')}  |  {p.get('program_name','')}  "
                f"|  Grant Narrative  |  {today}"
            )
            fr.font.size = Pt(8); fr.font.color.rgb = GRAY; fr.font.name = "Arial"

            docx_buf = io.BytesIO()
            doc.save(docx_buf)
            docx_buf.seek(0)

            st.download_button(
                "⬇ Download as Word doc",
                data=docx_buf.getvalue(),
                file_name=f"grant_narrative_{p.get('funder_name','').replace(' ','_')}_{today}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

        except ImportError:
            st.info("python-docx not available — use the text file download above.")

    if st.button("🔄 Start over with new program data"):
        st.session_state.narrative_generated = False
        st.session_state.narrative_sections  = {}
        st.rerun()
