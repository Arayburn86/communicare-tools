"""
Logic Model Builder — Communicare Alliance
==========================================
Upload any file (PDF, Word, Excel, text) and Claude extracts
your program data and builds a formatted logic model Word doc.

Install:  pip install streamlit anthropic python-docx PyPDF2 openpyxl
Run:      python -m streamlit run logic_model_builder.py

Needs your API key set first:
  set ANTHROPIC_API_KEY=sk-ant-your-key-here
"""

import streamlit as st
import anthropic
import json
import os
import re
import datetime
import subprocess
import tempfile
import io

# ─────────────────────────────────────────────────────────────────────────────
# Page config
# ─────────────────────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="Logic Model Builder — Communicare Alliance",
    page_icon="📋",
    layout="wide",
)

# ─────────────────────────────────────────────────────────────────────────────
# Styling
# ─────────────────────────────────────────────────────────────────────────────

st.markdown("""
<style>
  .main { background: #F8FAFB; }
  .block-container { padding-top: 1.5rem; padding-bottom: 2rem; }

  h1 { color: #1A6E6E !important; }
  h2 { color: #1A6E6E !important; font-size: 1.1rem !important; }
  h3 { color: #1F2937 !important; font-size: 1rem !important; }

  .stButton > button {
    background: #1A6E6E !important;
    color: white !important;
    border: none !important;
    border-radius: 8px !important;
    padding: 0.5rem 1.5rem !important;
    font-weight: 600 !important;
  }
  .stButton > button:hover {
    background: #0F6E56 !important;
  }

  .section-card {
    background: white;
    border-radius: 10px;
    padding: 16px 20px;
    margin-bottom: 12px;
    box-shadow: 0 1px 3px rgba(0,0,0,0.06);
    border-left: 4px solid #1A6E6E;
  }
  .section-card.blue   { border-left-color: #1E40AF; }
  .section-card.purple { border-left-color: #3C3489; }
  .section-card.green  { border-left-color: #065F46; }
  .section-card.amber  { border-left-color: #92400E; }
  .section-card.coral  { border-left-color: #993C1D; }
  .section-card.gray   { border-left-color: #5F5E5A; }

  .kpi-box {
    background: #E1F5EE;
    border-radius: 8px;
    padding: 12px;
    text-align: center;
    margin-bottom: 8px;
  }
  .kpi-val  { font-size: 2rem; font-weight: 700; color: #1A6E6E; }
  .kpi-lbl  { font-size: 0.75rem; color: #6B7280; }

  .upload-hint {
    background: #E1F5EE;
    border-radius: 8px;
    padding: 10px 14px;
    font-size: 0.85rem;
    color: #085041;
    margin-bottom: 12px;
  }
  .warning-box {
    background: #FEF3C7;
    border-left: 3px solid #BA7517;
    border-radius: 0 8px 8px 0;
    padding: 10px 14px;
    font-size: 0.85rem;
    color: #92400E;
    margin-bottom: 12px;
  }
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# Session state defaults
# ─────────────────────────────────────────────────────────────────────────────

FIELDS = [
    "inputs", "activities", "outputs",
    "short_term_outcomes", "medium_term_outcomes", "long_term_outcomes",
    "assumptions", "external_factors",
]

FIELD_META = {
    "inputs":               ("Inputs",               "teal",   "Resources you invest — staff, funding, facilities, partnerships"),
    "activities":           ("Activities",            "blue",   "What you do — services, training, outreach"),
    "outputs":              ("Outputs",               "purple", "Countable direct products — sessions, participants, events"),
    "short_term_outcomes":  ("Short-Term Outcomes",   "green",  "Changes 0–6 months after program participation"),
    "medium_term_outcomes": ("Medium-Term Outcomes",  "amber",  "Changes 6–24 months after participation"),
    "long_term_outcomes":   ("Long-Term Outcomes",    "coral",  "Broader changes 2–5 years out"),
    "assumptions":          ("Assumptions",           "gray",   "Conditions that must hold for the theory of change to work"),
    "external_factors":     ("External Factors",      "gray",   "Conditions outside your control that could affect outcomes"),
}

for field in FIELDS:
    if field not in st.session_state:
        st.session_state[field] = []

if "program_name"  not in st.session_state: st.session_state.program_name  = ""
if "org_name"      not in st.session_state: st.session_state.org_name      = "Communicare Alliance"
if "org_location"  not in st.session_state: st.session_state.org_location  = "Woonsocket, RI"
if "funder"        not in st.session_state: st.session_state.funder        = ""
if "period"        not in st.session_state: st.session_state.period        = ""
if "extracted"     not in st.session_state: st.session_state.extracted     = False
if "docx_bytes"    not in st.session_state: st.session_state.docx_bytes    = None

# ─────────────────────────────────────────────────────────────────────────────
# Text extraction from uploaded files
# ─────────────────────────────────────────────────────────────────────────────

def extract_text(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    raw  = uploaded_file.read()

    if name.endswith(".pdf"):
        try:
            import PyPDF2, io
            reader = PyPDF2.PdfReader(io.BytesIO(raw))
            return "\n\n".join(
                p.extract_text() for p in reader.pages
                if p.extract_text()
            )
        except ImportError:
            st.error("Run:  pip install PyPDF2")
            return ""

    elif name.endswith(".docx"):
        try:
            from docx import Document
            import io
            doc = Document(io.BytesIO(raw))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            st.error("Run:  pip install python-docx")
            return ""

    elif name.endswith((".xlsx", ".xls")):
        try:
            import pandas as pd, io
            sheets = pd.read_excel(io.BytesIO(raw), sheet_name=None)
            parts = []
            for sheet_name, df in sheets.items():
                parts.append(f"Sheet: {sheet_name}")
                parts.append(df.to_string(index=False))
            return "\n\n".join(parts)
        except ImportError:
            st.error("Run:  pip install openpyxl pandas")
            return ""

    elif name.endswith(".txt"):
        return raw.decode("utf-8", errors="ignore")

    else:
        try:
            return raw.decode("utf-8", errors="ignore")
        except Exception:
            return ""


# ─────────────────────────────────────────────────────────────────────────────
# Claude extraction
# ─────────────────────────────────────────────────────────────────────────────

def extract_logic_model(text: str, api_key: str) -> dict:
    client = anthropic.Anthropic(api_key=api_key)

    prompt = f"""You are an expert nonprofit program evaluator. Read this document and
extract a complete logic model for the program described.

Return ONLY a valid JSON object with these exact keys:
{{
  "program_name": "name of the program",
  "org_name": "name of the organization",
  "funder": "funder name if mentioned, else empty string",
  "period": "grant period if mentioned, else empty string",
  "inputs": ["list of resources invested — staff, funding, facilities, partnerships, curricula"],
  "activities": ["list of specific program activities — what staff actually do"],
  "outputs": ["list of direct countable products — sessions, participants, events, materials"],
  "short_term_outcomes": ["changes 0-6 months — knowledge, skills, attitudes, early behaviour change"],
  "medium_term_outcomes": ["changes 6-24 months — behaviour change, improved conditions"],
  "long_term_outcomes": ["changes 2-5 years — systemic change, population-level impact"],
  "assumptions": ["conditions that must hold for theory of change to work"],
  "external_factors": ["outside conditions that could affect outcomes"]
}}

Rules:
- Extract only what is actually in the document — do not invent items
- Each list item should be a clear concise statement (1-2 sentences max)
- Aim for 5-12 items per list where the document supports it
- If a section cannot be determined from the document, return an empty list []
- Return ONLY the JSON object, no other text

DOCUMENT:
{text[:12000]}"""

    msg = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=3000,
        messages=[{"role": "user", "content": prompt}],
    )
    raw = msg.content[0].text.strip()
    raw = re.sub(r"^```json\s*", "", raw)
    raw = re.sub(r"```\s*$", "", raw)
    return json.loads(raw)


# ─────────────────────────────────────────────────────────────────────────────
# Word doc builder (via docx-js)
# ─────────────────────────────────────────────────────────────────────────────

def build_docx(program: dict) -> bytes:
    today   = datetime.date.today().strftime("%B %d, %Y")
    payload = json.dumps({**program, "today": today}, ensure_ascii=False)

    js = r"""
const { Document, Packer, Paragraph, TextRun, Table, TableRow, TableCell,
        AlignmentType, BorderStyle, WidthType, ShadingType,
        LevelFormat, VerticalAlign } = require('docx');
const fs = require('fs');

const D = """ + payload + r""";
const OUT = D.output_path;

const TEAL="1A6E6E",TEAL_P="E1F5EE",TEAL_M="9FE1CB";
const WHITE="FFFFFF",GRAY_L="F5F5F5",GRAY_M="D1D5DB",DARK="1F2937";
const AMBER="92400E",AMBER_L="FEF3C7",BLUE="1E40AF",BLUE_L="DBEAFE";
const GREEN="065F46",GREEN_L="D1FAE5",PURPLE="3C3489",PURPLE_L="EEEDFE";
const CORAL="993C1D",CORAL_L="FAECE7";

const COLUMNS=[
  {key:"inputs",            label:"INPUTS",               hbg:TEAL,   bbg:TEAL_P,   tc:"0F6E56"},
  {key:"activities",        label:"ACTIVITIES",            hbg:BLUE,   bbg:BLUE_L,   tc:BLUE},
  {key:"outputs",           label:"OUTPUTS",               hbg:PURPLE, bbg:PURPLE_L, tc:PURPLE},
  {key:"short_term_outcomes",label:"SHORT-TERM\nOUTCOMES", hbg:GREEN,  bbg:GREEN_L,  tc:GREEN},
  {key:"medium_term_outcomes",label:"MEDIUM-TERM\nOUTCOMES",hbg:AMBER, bbg:AMBER_L,  tc:AMBER},
  {key:"long_term_outcomes",label:"LONG-TERM\nOUTCOMES",   hbg:CORAL,  bbg:CORAL_L,  tc:CORAL},
];

const TOTAL_W=14040;
const COL_W=Math.floor(TOTAL_W/COLUMNS.length);
const colWidths=COLUMNS.map((_,i)=>i<COLUMNS.length-1?COL_W:TOTAL_W-COL_W*(COLUMNS.length-1));

const bdr=(c=GRAY_M)=>({style:BorderStyle.SINGLE,size:4,color:c});
const bdrs=c=>({top:bdr(c),bottom:bdr(c),left:bdr(c),right:bdr(c)});
const shd=h=>({fill:h,type:ShadingType.CLEAR});
const mg=(t=80,b=80,l=100,r=100)=>({top:t,bottom:b,left:l,right:r});

function tr(text,opts={}){
  return new TextRun({text,font:"Arial",size:opts.size||18,bold:opts.bold||false,
    color:opts.color||DARK,italics:opts.italic||false});
}
function p(text,opts={}){
  return new Paragraph({children:[tr(text,opts)],
    spacing:{before:opts.before||0,after:opts.after||60},
    alignment:opts.align||AlignmentType.LEFT});
}
function bullet(text,color=DARK){
  return new Paragraph({numbering:{reference:"bullets",level:0},
    spacing:{before:0,after:40},children:[tr(text,{size:17,color})]});
}

const logicTable=new Table({
  width:{size:TOTAL_W,type:WidthType.DXA},columnWidths:colWidths,
  rows:[
    // Title row
    new TableRow({children:[new TableCell({columnSpan:COLUMNS.length,
      borders:bdrs(TEAL),width:{size:TOTAL_W,type:WidthType.DXA},
      shading:shd(TEAL),margins:mg(120,120,160,160),verticalAlign:VerticalAlign.CENTER,
      children:[
        new Paragraph({alignment:AlignmentType.CENTER,spacing:{before:0,after:40},
          children:[tr(D.program_name||"Program Logic Model",{bold:true,size:30,color:WHITE})]}),
        new Paragraph({alignment:AlignmentType.CENTER,spacing:{before:0,after:0},
          children:[
            tr((D.org_name||"")+(D.org_location?" \u00b7 "+D.org_location:""),{size:20,color:TEAL_M}),
            tr(D.funder?" \u00b7 "+D.funder:"",{size:20,color:TEAL_M}),
            tr(D.period?" \u00b7 "+D.period:"",{size:20,color:TEAL_M}),
          ]}),
      ]})]
    }),
    // Theory row
    new TableRow({height:{value:280,rule:"atLeast"},children:[
      ...[{label:"RESOURCES",span:1,bg:TEAL_P,tc:TEAL},
          {label:"PROGRAM IMPLEMENTATION",span:2,bg:BLUE_L,tc:BLUE},
          {label:"\u2190\u2014\u2014\u2014\u2014 RESULTS \u2014\u2014\u2014\u2014\u2192",span:3,bg:GREEN_L,tc:GREEN}]
        .map((g,idx)=>{
          const startCol=idx===0?0:idx===1?1:3;
          const w=colWidths.slice(startCol,startCol+g.span).reduce((a,b)=>a+b,0);
          return new TableCell({columnSpan:g.span,borders:bdrs(GRAY_M),
            width:{size:w,type:WidthType.DXA},shading:shd(g.bg),margins:mg(60,60,120,120),
            verticalAlign:VerticalAlign.CENTER,
            children:[new Paragraph({alignment:AlignmentType.CENTER,spacing:{before:0,after:0},
              children:[tr(g.label,{bold:true,size:16,color:g.tc})]})]});
        })
    ]}),
    // Header row
    new TableRow({tableHeader:true,height:{value:500,rule:"atLeast"},
      children:COLUMNS.map((col,i)=>new TableCell({
        borders:bdrs(col.hbg),width:{size:colWidths[i],type:WidthType.DXA},
        shading:shd(col.hbg),margins:mg(100,100,120,120),verticalAlign:VerticalAlign.CENTER,
        children:col.label.split("\n").map((line,li)=>new Paragraph({
          alignment:AlignmentType.CENTER,
          spacing:{before:0,after:li===0&&col.label.includes("\n")?20:0},
          children:[tr(line,{bold:true,size:18,color:WHITE})]}))
      }))
    }),
    // Data row
    new TableRow({children:COLUMNS.map((col,i)=>new TableCell({
      borders:bdrs(GRAY_M),width:{size:colWidths[i],type:WidthType.DXA},
      shading:shd(col.bbg),margins:mg(100,100,120,120),verticalAlign:VerticalAlign.TOP,
      children:(D[col.key]||[]).map(item=>bullet(item,col.tc))
    }))}),
  ],
});

function smallTable(label,items,hbg,bbg,tc){
  const W=9360;
  return new Table({width:{size:W,type:WidthType.DXA},columnWidths:[W],rows:[
    new TableRow({children:[new TableCell({borders:bdrs(hbg),width:{size:W,type:WidthType.DXA},
      shading:shd(hbg),margins:mg(80,80,120,120),
      children:[p(label,{bold:true,size:18,color:WHITE})]})]}),
    new TableRow({children:[new TableCell({borders:bdrs(GRAY_M),width:{size:W,type:WidthType.DXA},
      shading:shd(bbg),margins:mg(100,100,120,120),verticalAlign:VerticalAlign.TOP,
      children:(items||[]).map(item=>bullet(item,tc))})]})
  ]});
}

const doc=new Document({
  styles:{default:{document:{run:{font:"Arial",size:18}}}},
  numbering:{config:[{reference:"bullets",levels:[{
    level:0,format:LevelFormat.BULLET,text:"\u2022",alignment:AlignmentType.LEFT,
    style:{paragraph:{indent:{left:360,hanging:260}}}}]}]},
  sections:[
    {properties:{page:{size:{width:12240,height:15840,orientation:"landscape"},
      margin:{top:720,right:900,bottom:720,left:900}}},
     children:[logicTable,
       new Paragraph({spacing:{before:120,after:0},alignment:AlignmentType.CENTER,
         children:[tr((D.org_name||"")+" \u00b7 Logic Model \u00b7 "+(D.program_name||"")+" \u00b7 "+D.today,
           {size:14,color:"AAAAAA",italic:true})]})]},
    {properties:{page:{size:{width:12240,height:15840},
      margin:{top:1080,right:1080,bottom:1080,left:1080}}},
     children:[
       p((D.program_name||"Program")+" \u2014 Logic Model Supporting Notes",
         {bold:true,size:24,color:TEAL,after:120}),
       p("Generated: "+D.today,{size:16,color:"888888",italic:true,after:200}),
       smallTable("ASSUMPTIONS",D.assumptions,TEAL,TEAL_P,"0F6E56"),
       new Paragraph({spacing:{before:200,after:0},children:[]}),
       smallTable("EXTERNAL FACTORS",D.external_factors,"5F5E5A",GRAY_L,"444441"),
       new Paragraph({spacing:{before:240,after:0},alignment:AlignmentType.CENTER,
         children:[tr((D.org_name||"")+" \u00b7 Generated "+D.today,
           {size:14,color:"AAAAAA",italic:true})]}),
     ]},
  ],
});

Packer.toBuffer(doc).then(buf=>{
  fs.writeFileSync(OUT,buf);
  console.log("saved:"+OUT);
});
"""

    with tempfile.TemporaryDirectory() as tmpdir:
        js_path   = os.path.join(tmpdir, "build.js")
        docx_path = os.path.join(tmpdir, "logic_model.docx")
        program["output_path"] = docx_path

        # Re-encode payload with output_path
        js = js.replace(
            "const D = " + payload,
            "const D = " + json.dumps({**program, "today": today}, ensure_ascii=False)
        )
        with open(js_path, "w", encoding="utf-8") as f:
            f.write(js)

        result = subprocess.run(
            ["node", js_path],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode != 0:
            raise RuntimeError(result.stderr[:400])

        with open(docx_path, "rb") as f:
            return f.read()


# ─────────────────────────────────────────────────────────────────────────────
# Editable list component
# ─────────────────────────────────────────────────────────────────────────────

def editable_list(field: str, label: str, color: str, hint: str):
    items = st.session_state[field]
    n     = len(items)

    st.markdown(
        f'<div class="section-card {color}">'
        f'<strong>{label}</strong> <span style="color:#9CA3AF;font-size:0.8rem">— {hint}</span>'
        f'</div>',
        unsafe_allow_html=True
    )

    to_delete = []
    for i, item in enumerate(items):
        col1, col2 = st.columns([11, 1])
        with col1:
            new_val = st.text_input(
                f"{label} item {i+1}",
                value=item,
                key=f"{field}_{i}",
                label_visibility="collapsed",
            )
            st.session_state[field][i] = new_val
        with col2:
            if st.button("✕", key=f"del_{field}_{i}", help="Remove"):
                to_delete.append(i)

    for idx in reversed(to_delete):
        st.session_state[field].pop(idx)
        st.rerun()

    new_item = st.text_input(
        f"Add to {label}",
        key=f"new_{field}",
        placeholder=f"+ Add a {label.lower()[:-1] if label.endswith('s') else label.lower()}...",
        label_visibility="collapsed",
    )
    if new_item:
        st.session_state[field].append(new_item)
        st.rerun()

    st.caption(f"{n} item{'s' if n != 1 else ''}")
    st.markdown("---")


# ─────────────────────────────────────────────────────────────────────────────
# Main app
# ─────────────────────────────────────────────────────────────────────────────

def main():
    # Header
    st.title("📋 Logic Model Builder")
    st.caption("Upload a document — Claude extracts your program data and builds a formatted Word logic model.")

    # API key check
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        st.markdown(
            '<div class="warning-box">⚠ API key not set. '
            'Run <code>set ANTHROPIC_API_KEY=sk-ant-...</code> in Command Prompt '
            'then restart Streamlit.</div>',
            unsafe_allow_html=True
        )

    # ── Sidebar ────────────────────────────────────────────────────────────────
    with st.sidebar:
        st.markdown("### Program details")
        st.session_state.org_name     = st.text_input("Organization", st.session_state.org_name)
        st.session_state.org_location = st.text_input("Location",     st.session_state.org_location)
        st.session_state.program_name = st.text_input("Program name", st.session_state.program_name)
        st.session_state.funder       = st.text_input("Funder",       st.session_state.funder)
        st.session_state.period       = st.text_input("Grant period",  st.session_state.period)

        st.markdown("---")
        st.markdown("### How to use")
        st.markdown("""
1. Upload your file below
2. Click **Extract with Claude**
3. Review and edit each section
4. Click **Generate Logic Model**
5. Download your Word doc
        """)

        st.markdown("---")
        st.markdown("### Item counts")
        for field in FIELDS:
            label, _, _ = FIELD_META[field]
            n = len(st.session_state[field])
            color = "🟢" if n >= 3 else ("🟡" if n >= 1 else "⚪")
            st.caption(f"{color} {label}: {n}")

        if st.button("🗑 Clear all data", use_container_width=True):
            for field in FIELDS:
                st.session_state[field] = []
            st.session_state.extracted  = False
            st.session_state.docx_bytes = None
            st.rerun()

    # ── Upload + extract ───────────────────────────────────────────────────────
    st.markdown("## Step 1 — Upload your document")
    st.markdown(
        '<div class="upload-hint">Accepts: grant narrative (PDF or Word), '
        'program description, tracking spreadsheet (Excel), or any text document. '
        'The more detail your file contains, the better the logic model.</div>',
        unsafe_allow_html=True
    )

    uploaded = st.file_uploader(
        "Drop your file here",
        type=["pdf", "docx", "doc", "xlsx", "xls", "txt"],
        label_visibility="collapsed",
    )

    if uploaded:
        col1, col2 = st.columns([3, 1])
        with col1:
            st.success(f"✅ Uploaded: **{uploaded.name}**  ({uploaded.size:,} bytes)")
        with col2:
            extract_btn = st.button(
                "✨ Extract with Claude",
                disabled=not api_key,
                use_container_width=True,
            )

        if extract_btn:
            with st.spinner("Reading your file and extracting program data..."):
                try:
                    text = extract_text(uploaded)
                    if not text.strip():
                        st.error("Could not read text from this file. Try saving as .txt or .docx.")
                    else:
                        st.info(f"Extracted {len(text):,} characters. Sending to Claude...")
                        result = extract_logic_model(text, api_key)

                        # Load extracted data into session state
                        for field in FIELDS:
                            if field in result and result[field]:
                                st.session_state[field] = result[field]

                        # Update program details if found
                        if result.get("program_name"):
                            st.session_state.program_name = result["program_name"]
                        if result.get("org_name") and result["org_name"] != "Communicare Alliance":
                            st.session_state.org_name = result["org_name"]
                        if result.get("funder"):
                            st.session_state.funder = result["funder"]
                        if result.get("period"):
                            st.session_state.period = result["period"]

                        st.session_state.extracted  = True
                        st.session_state.docx_bytes = None

                        total = sum(len(st.session_state[f]) for f in FIELDS)
                        st.success(f"✅ Extracted {total} items across {len(FIELDS)} sections. Review and edit below.")
                        st.rerun()

                except json.JSONDecodeError as e:
                    st.error(f"Claude returned unexpected output. Try again or add items manually. ({e})")
                except Exception as e:
                    st.error(f"Extraction failed: {e}")

    st.markdown("---")

    # ── Manual / edit section ──────────────────────────────────────────────────
    st.markdown("## Step 2 — Review and edit")

    if not st.session_state.extracted:
        st.markdown(
            '<div class="upload-hint">Upload a file above and click '
            '<strong>Extract with Claude</strong>, or add items manually below.</div>',
            unsafe_allow_html=True
        )

    # Counts banner
    total_items = sum(len(st.session_state[f]) for f in FIELDS)
    if total_items > 0:
        cols = st.columns(len(FIELDS))
        for i, field in enumerate(FIELDS):
            label, _, _ = FIELD_META[field]
            with cols[i]:
                st.markdown(
                    f'<div class="kpi-box">'
                    f'<div class="kpi-val">{len(st.session_state[field])}</div>'
                    f'<div class="kpi-lbl">{label.split()[0]}</div>'
                    f'</div>',
                    unsafe_allow_html=True
                )

    # Editable sections in two columns
    col_left, col_right = st.columns(2)

    with col_left:
        for field in ["inputs", "activities", "outputs", "assumptions"]:
            label, color, hint = FIELD_META[field]
            editable_list(field, label, color, hint)

    with col_right:
        for field in ["short_term_outcomes", "medium_term_outcomes",
                      "long_term_outcomes", "external_factors"]:
            label, color, hint = FIELD_META[field]
            editable_list(field, label, color, hint)

    # ── Generate ───────────────────────────────────────────────────────────────
    st.markdown("## Step 3 — Generate")

    has_content = any(len(st.session_state[f]) > 0 for f in FIELDS)

    if not has_content:
        st.info("Add at least one item above to generate the logic model.")
    else:
        gen_col, dl_col = st.columns([2, 1])

        with gen_col:
            if st.button("📄 Generate Logic Model Word Doc", use_container_width=True):
                program = {
                    "org_name":      st.session_state.org_name,
                    "org_location":  st.session_state.org_location,
                    "program_name":  st.session_state.program_name,
                    "funder":        st.session_state.funder,
                    "period":        st.session_state.period,
                    "preparer":      "Program Development & Grants Manager",
                }
                for field in FIELDS:
                    program[field] = st.session_state[field]

                with st.spinner("Building your Word document..."):
                    try:
                        docx_bytes = build_docx(program)
                        st.session_state.docx_bytes = docx_bytes
                        st.success("✅ Logic model ready to download!")
                    except FileNotFoundError:
                        st.error(
                            "Node.js is not installed. Download it free at nodejs.org "
                            "then restart Streamlit."
                        )
                    except Exception as e:
                        st.error(f"Build failed: {e}")

        with dl_col:
            if st.session_state.docx_bytes:
                prog_slug = (st.session_state.program_name or "logic_model")\
                    .lower().replace(" ", "_")[:30]
                fname = f"{prog_slug}_logic_model.docx"
                st.download_button(
                    label="⬇ Download .docx",
                    data=st.session_state.docx_bytes,
                    file_name=fname,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )

    # Footer
    st.markdown("---")
    st.caption(
        "Communicare Alliance — Logic Model Builder  ·  "
        "AI-assisted extraction — always review before attaching to a grant application"
    )


if __name__ == "__main__":
    main()
