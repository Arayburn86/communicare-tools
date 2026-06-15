"""
Grant Compliance Checker
=========================
Reads your RFP (PDF or Word) and draft narrative (Word or text),
uses the Claude API to extract every requirement from the RFP, then
checks each one against your draft and produces a colour-coded report.

Install:  pip install anthropic python-docx PyPDF2
Run:      python compliance_checker.py

Or in Google Colab — see COLAB INSTRUCTIONS at the bottom of this file.

Outputs:
  - Terminal summary (pass / partial / missing counts)
  - compliance_report.docx  — detailed Word report with all findings
  - compliance_report.html  — shareable HTML version (open in any browser)
"""

import anthropic
import json
import os
import re
import datetime
import sys

# ─────────────────────────────────────────────────────────────────────────────
# ✏️  CONFIGURATION
# ─────────────────────────────────────────────────────────────────────────────

RFP_FILE       = "rfp.pdf"          # PDF or .docx — your RFP / NOFO
DRAFT_FILE     = "draft.docx"       # .docx or .txt — your draft narrative
REPORT_DOCX    = "compliance_report.docx"
REPORT_HTML    = "compliance_report.html"

ORG_NAME       = "Communicare Alliance"
FUNDER_NAME    = "Rhode Island Foundation"  # Used in report header

# ─────────────────────────────────────────────────────────────────────────────
# Text extraction
# ─────────────────────────────────────────────────────────────────────────────

def extract_text_from_pdf(path: str) -> str:
    try:
        import PyPDF2
        text = []
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text.append(t)
        return "\n\n".join(text)
    except ImportError:
        raise ImportError("Run:  pip install PyPDF2")


def extract_text_from_docx(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except ImportError:
        raise ImportError("Run:  pip install python-docx")


def extract_text(path: str) -> str:
    """Auto-detect file type and extract text."""
    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: {path}")
    ext = path.lower().split(".")[-1]
    if ext == "pdf":
        return extract_text_from_pdf(path)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(path)
    elif ext == "txt":
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {ext}. Use PDF, DOCX, or TXT.")


# ─────────────────────────────────────────────────────────────────────────────
# Claude API calls
# ─────────────────────────────────────────────────────────────────────────────

def extract_requirements(client: anthropic.Anthropic, rfp_text: str) -> list[dict]:
    """
    Ask Claude to extract every checkable requirement from the RFP.
    Returns a list of dicts: {id, category, requirement, mandatory, source_quote}
    """
    print("  Extracting requirements from RFP...")

    prompt = f"""You are a grant compliance specialist. Read this RFP/NOFO carefully and 
extract EVERY distinct requirement an applicant must address in their narrative.

Include:
- Mandatory elements (must, shall, required, applicants must demonstrate)
- Narrative content requirements (describe, explain, provide, address)
- Page/word limits and formatting rules
- Eligibility requirements that should be demonstrated
- Evaluation criteria that should be addressed
- Required attachments or appendices mentioned
- Budget narrative requirements

For each requirement output a JSON object with:
  "id": sequential number (1, 2, 3...)
  "category": one of: Eligibility | Program Design | Evaluation | Budget | 
               Organizational Capacity | Formatting | Attachments | Other
  "requirement": the requirement stated clearly in plain language (1-2 sentences)
  "mandatory": true if explicitly required, false if recommended/preferred
  "source_quote": the exact phrase from the RFP that states this requirement 
                  (max 20 words)

Return ONLY a valid JSON array. No prose, no markdown, no code fences.

RFP TEXT:
{rfp_text[:12000]}
"""

    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4000,
        messages=[{"role": "user", "content": prompt}],
    )

    raw = message.content[0].text.strip()
    # Strip any accidental markdown fences
    raw = re.sub(r"^```json\s*", "", raw)
    raw = re.sub(r"```\s*$", "", raw)

    try:
        requirements = json.loads(raw)
        print(f"  Found {len(requirements)} requirements")
        return requirements
    except json.JSONDecodeError:
        # Try to salvage partial JSON
        match = re.search(r'\[.*\]', raw, re.DOTALL)
        if match:
            return json.loads(match.group(0))
        raise ValueError(f"Could not parse requirements JSON. Raw output:\n{raw[:500]}")


def check_requirement(client: anthropic.Anthropic,
                       requirement: dict,
                       draft_text: str) -> dict:
    """
    Check a single requirement against the draft narrative.
    Returns the requirement dict enriched with:
      status: "Pass" | "Partial" | "Missing"
      confidence: 1-5
      finding: explanation of what was found or missing
      recommendation: specific suggested improvement (if not Pass)
      draft_excerpt: relevant text from draft (if found)
    """
    prompt = f"""You are reviewing a grant narrative draft for compliance with a specific requirement.

REQUIREMENT #{requirement['id']} ({requirement['category']}):
{requirement['requirement']}

Source from RFP: "{requirement.get('source_quote', '')}"

DRAFT NARRATIVE (first 8000 chars):
{draft_text[:8000]}

Assess whether the draft addresses this requirement. Return ONLY a JSON object with:
  "status": "Pass" if fully addressed, "Partial" if mentioned but needs strengthening,
            "Missing" if not addressed at all
  "confidence": integer 1-5 (how confident you are in this assessment)
  "finding": 1-2 sentence explanation of what you found in the draft
  "recommendation": if status is Partial or Missing, give a specific 1-2 sentence 
                    suggestion for what to add or strengthen. If Pass, write "None needed."
  "draft_excerpt": if found, paste the most relevant 1-2 sentences from the draft 
                   that address this requirement. If not found, write "Not found."

Return ONLY valid JSON. No prose, no markdown fences."""

    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=600,
        messages=[{"role": "user", "content": prompt}],
    )

    raw = message.content[0].text.strip()
    raw = re.sub(r"^```json\s*", "", raw)
    raw = re.sub(r"```\s*$", "", raw)

    try:
        result = json.loads(raw)
        return {**requirement, **result}
    except json.JSONDecodeError:
        return {
            **requirement,
            "status": "Unknown",
            "confidence": 1,
            "finding": "Could not parse Claude response",
            "recommendation": "Review manually",
            "draft_excerpt": "N/A",
        }


def run_compliance_check(rfp_text: str,
                          draft_text: str,
                          api_key: str) -> list[dict]:
    """Run the full compliance check and return enriched requirements list."""
    client = anthropic.Anthropic(api_key=api_key)

    requirements = extract_requirements(client, rfp_text)

    print(f"\n  Checking draft against {len(requirements)} requirements...")
    results = []
    for i, req in enumerate(requirements, 1):
        print(f"    [{i:2}/{len(requirements)}] #{req['id']} {req['category']}: "
              f"{req['requirement'][:60]}...")
        result = check_requirement(client, req, draft_text)
        results.append(result)

    return results


# ─────────────────────────────────────────────────────────────────────────────
# Word report builder (JS via Node)
# ─────────────────────────────────────────────────────────────────────────────

def build_word_report(results: list[dict], output_path: str,
                       rfp_file: str, draft_file: str):
    """Generate compliance_report.docx using docx-js via Node."""
    import subprocess, tempfile

    today     = datetime.date.today().strftime("%B %d, %Y")
    total     = len(results)
    passed    = sum(1 for r in results if r.get("status") == "Pass")
    partial   = sum(1 for r in results if r.get("status") == "Partial")
    missing   = sum(1 for r in results if r.get("status") == "Missing")
    mandatory_missing = sum(1 for r in results
                            if r.get("status") == "Missing" and r.get("mandatory", True))
    score     = round(passed / total * 100) if total else 0

    # Serialize results for JS
    data_json = json.dumps({
        "org": ORG_NAME,
        "funder": FUNDER_NAME,
        "rfp_file": rfp_file,
        "draft_file": draft_file,
        "today": today,
        "total": total,
        "passed": passed,
        "partial": partial,
        "missing": missing,
        "mandatory_missing": mandatory_missing,
        "score": score,
        "results": results,
    })

    js = r"""
const { Document, Packer, Paragraph, TextRun, Table, TableRow, TableCell,
        AlignmentType, BorderStyle, WidthType, ShadingType, HeadingLevel, LevelFormat } = require('docx');
const fs = require('fs');

const DATA = """ + data_json + r""";
const outputPath = process.argv[2];

const TEAL    = "1A6E6E";
const TEAL_P  = "E1F5EE";
const WHITE   = "FFFFFF";
const DARK    = "1F2937";
const GRAY_L  = "F5F5F5";
const GRAY_M  = "D1D5DB";
const GREEN_L = "D1FAE5";
const GREEN_T = "065F46";
const AMBER_L = "FEF3C7";
const AMBER_T = "92400E";
const RED_L   = "FDECEA";
const RED_T   = "C0392B";

const border = (color = GRAY_M) => ({ style: BorderStyle.SINGLE, size: 4, color });
const borders = (c) => ({ top: border(c), bottom: border(c), left: border(c), right: border(c) });
const fill = (hex) => ({ fill: hex, type: ShadingType.CLEAR });

const STATUS_COLORS = {
  "Pass":    { bg: GREEN_L, text: GREEN_T, label: "PASS" },
  "Partial": { bg: AMBER_L, text: AMBER_T, label: "PARTIAL" },
  "Missing": { bg: RED_L,   text: RED_T,   label: "MISSING" },
  "Unknown": { bg: GRAY_L,  text: DARK,    label: "UNKNOWN" },
};

const CAT_COLORS = {
  "Eligibility":              "534AB7",
  "Program Design":           "1A6E6E",
  "Evaluation":               "0C447C",
  "Budget":                   "854F0B",
  "Organizational Capacity":  "3B6D11",
  "Formatting":               "5F5E5A",
  "Attachments":              "993C1D",
  "Other":                    "444441",
};

function para(text, opts = {}) {
  return new Paragraph({
    spacing: { before: opts.spaceBefore || 80, after: opts.spaceAfter || 80 },
    alignment: opts.align || AlignmentType.LEFT,
    children: [new TextRun({
      text,
      bold: opts.bold || false,
      size: opts.size || 20,
      color: opts.color || DARK,
      font: "Arial",
      italics: opts.italic || false,
    })],
  });
}

function divider() {
  return new Paragraph({
    spacing: { before: 120, after: 0 },
    border: { bottom: { style: BorderStyle.SINGLE, size: 4, color: "9FE1CB", space: 1 } },
    children: [],
  });
}

function kpiTable(data) {
  const kpis = [
    { label: "Total requirements", value: String(data.total),  bg: TEAL_P,  text: TEAL },
    { label: "Passing",            value: String(data.passed),  bg: GREEN_L, text: GREEN_T },
    { label: "Partial",            value: String(data.partial), bg: AMBER_L, text: AMBER_T },
    { label: "Missing",            value: String(data.missing), bg: RED_L,   text: RED_T },
    { label: "Compliance score",   value: data.score + "%",     bg: TEAL_P,  text: TEAL },
  ];
  const w = 1800;
  return new Table({
    width: { size: 9000, type: WidthType.DXA },
    columnWidths: kpis.map(() => w),
    rows: [
      new TableRow({ children: kpis.map(k => new TableCell({
        borders: borders(GRAY_M),
        width: { size: w, type: WidthType.DXA },
        shading: fill(k.bg),
        margins: { top: 120, bottom: 120, left: 120, right: 120 },
        children: [
          new Paragraph({ alignment: AlignmentType.CENTER, spacing: { before: 0, after: 40 },
            children: [new TextRun({ text: k.value, bold: true, size: 36, color: k.text, font: "Arial" })] }),
          new Paragraph({ alignment: AlignmentType.CENTER, spacing: { before: 0, after: 0 },
            children: [new TextRun({ text: k.label, size: 16, color: k.text, font: "Arial" })] }),
        ],
      })) }),
    ],
  });
}

function requirementRow(req, i) {
  const sc = STATUS_COLORS[req.status] || STATUS_COLORS["Unknown"];
  const catColor = CAT_COLORS[req.category] || DARK;
  const alt = i % 2 === 0 ? GRAY_L : WHITE;

  // Column widths: #, Category, Requirement, Status, Finding, Recommendation
  const colWidths = [400, 1200, 2400, 700, 2200, 2460];

  return new TableRow({
    children: [
      // # 
      new TableCell({
        borders: borders(GRAY_M), width: { size: colWidths[0], type: WidthType.DXA },
        shading: fill(alt), margins: { top: 80, bottom: 80, left: 80, right: 80 },
        children: [new Paragraph({ alignment: AlignmentType.CENTER,
          children: [new TextRun({ text: String(req.id), bold: true, size: 18, color: DARK, font: "Arial" })] })],
      }),
      // Category
      new TableCell({
        borders: borders(GRAY_M), width: { size: colWidths[1], type: WidthType.DXA },
        shading: fill(alt), margins: { top: 80, bottom: 80, left: 100, right: 80 },
        children: [
          new Paragraph({ children: [new TextRun({ text: req.category || "", size: 17,
            bold: true, color: catColor, font: "Arial" })] }),
          new Paragraph({ children: [new TextRun({ text: req.mandatory ? "Required" : "Preferred",
            size: 15, italic: true, color: req.mandatory ? RED_T : "888888", font: "Arial" })] }),
        ],
      }),
      // Requirement
      new TableCell({
        borders: borders(GRAY_M), width: { size: colWidths[2], type: WidthType.DXA },
        shading: fill(alt), margins: { top: 80, bottom: 80, left: 100, right: 80 },
        children: [
          new Paragraph({ children: [new TextRun({ text: req.requirement || "", size: 18, color: DARK, font: "Arial" })] }),
          ...(req.source_quote ? [new Paragraph({ spacing: { before: 40, after: 0 },
            children: [new TextRun({ text: "\u201c" + req.source_quote + "\u201d",
              size: 16, italic: true, color: "888888", font: "Arial" })] })] : []),
        ],
      }),
      // Status badge
      new TableCell({
        borders: borders(GRAY_M), width: { size: colWidths[3], type: WidthType.DXA },
        shading: fill(sc.bg), margins: { top: 80, bottom: 80, left: 80, right: 80 },
        children: [new Paragraph({ alignment: AlignmentType.CENTER,
          children: [new TextRun({ text: sc.label, bold: true, size: 17, color: sc.text, font: "Arial" })] })],
      }),
      // Finding
      new TableCell({
        borders: borders(GRAY_M), width: { size: colWidths[4], type: WidthType.DXA },
        shading: fill(alt), margins: { top: 80, bottom: 80, left: 100, right: 80 },
        children: [
          new Paragraph({ children: [new TextRun({ text: req.finding || "", size: 17, color: DARK, font: "Arial" })] }),
          ...(req.draft_excerpt && req.draft_excerpt !== "Not found." ? [
            new Paragraph({ spacing: { before: 40, after: 0 },
              children: [new TextRun({ text: "Found: " + req.draft_excerpt,
                size: 15, italic: true, color: "555555", font: "Arial" })] })
          ] : []),
        ],
      }),
      // Recommendation
      new TableCell({
        borders: borders(GRAY_M), width: { size: colWidths[5], type: WidthType.DXA },
        shading: req.recommendation && req.recommendation !== "None needed." ? fill(sc.bg) : fill(alt),
        margins: { top: 80, bottom: 80, left: 100, right: 80 },
        children: [new Paragraph({ children: [new TextRun({
          text: req.recommendation || "None needed.",
          size: 17,
          color: req.recommendation && req.recommendation !== "None needed." ? sc.text : "888888",
          font: "Arial",
          italic: req.recommendation === "None needed.",
        })] })],
      }),
    ],
  });
}

// Group results by category
const categories = [...new Set(DATA.results.map(r => r.category))];

const children = [
  // Cover
  new Paragraph({ spacing: { before: 0, after: 80 },
    children: [new TextRun({ text: "Grant Compliance Report", bold: true, size: 44,
      color: TEAL, font: "Arial" })] }),
  new Paragraph({ spacing: { before: 0, after: 60 },
    children: [new TextRun({ text: DATA.org + "  \u2022  " + DATA.funder,
      size: 22, color: "888888", font: "Arial" })] }),
  new Paragraph({ spacing: { before: 0, after: 40 },
    children: [new TextRun({ text: "RFP: " + DATA.rfp_file + "  \u2022  Draft: " + DATA.draft_file,
      size: 18, color: "888888", font: "Arial", italics: true })] }),
  new Paragraph({ spacing: { before: 0, after: 200 },
    children: [new TextRun({ text: "Generated " + DATA.today,
      size: 18, color: "AAAAAA", font: "Arial", italics: true })] }),

  kpiTable(DATA),

  ...(DATA.mandatory_missing > 0 ? [
    new Paragraph({ spacing: { before: 200, after: 80 },
      children: [new TextRun({ text: "\u26a0\ufe0f  " + DATA.mandatory_missing +
        " mandatory requirement(s) are Missing — address these before submitting.",
        bold: true, size: 20, color: RED_T, font: "Arial" })] }),
  ] : [
    new Paragraph({ spacing: { before: 200, after: 80 },
      children: [new TextRun({ text: "\u2705  All mandatory requirements addressed. Review Partial items before submitting.",
        bold: true, size: 20, color: GREEN_T, font: "Arial" })] }),
  ]),

  divider(),

  // Results table
  new Paragraph({ spacing: { before: 200, after: 100 },
    children: [new TextRun({ text: "Detailed Findings", bold: true, size: 28, color: TEAL, font: "Arial" })] }),

  new Table({
    width: { size: 9360, type: WidthType.DXA },
    columnWidths: [400, 1200, 2400, 700, 2200, 2460],
    rows: [
      // Header
      new TableRow({ tableHeader: true, children: [
        "#", "Category", "Requirement", "Status", "What Was Found", "Recommendation"
      ].map((h, ci) => new TableCell({
        borders: borders(TEAL),
        width: { size: [400,1200,2400,700,2200,2460][ci], type: WidthType.DXA },
        shading: fill(TEAL),
        margins: { top: 100, bottom: 100, left: 100, right: 80 },
        children: [new Paragraph({ alignment: AlignmentType.CENTER,
          children: [new TextRun({ text: h, bold: true, size: 18, color: WHITE, font: "Arial" })] })],
      })) }),
      ...DATA.results.map((req, i) => requirementRow(req, i)),
    ],
  }),

  divider(),

  // Footer
  new Paragraph({ spacing: { before: 160, after: 0 }, alignment: AlignmentType.CENTER,
    children: [new TextRun({ text: DATA.org + "  \u2022  Compliance Report  \u2022  " + DATA.today,
      size: 16, color: "AAAAAA", font: "Arial" })] }),
];

const doc = new Document({
  styles: { default: { document: { run: { font: "Arial", size: 20 } } } },
  sections: [{
    properties: {
      page: { size: { width: 15840, height: 12240 },
               margin: { top: 1080, right: 900, bottom: 1080, left: 900 } }
    },
    children,
  }],
});

Packer.toBuffer(doc).then(buf => {
  fs.writeFileSync(outputPath, buf);
  console.log("saved:" + outputPath);
});
"""

    # Write JS to temp file and run it
    js_path = output_path.replace(".docx", "_builder.js")
    with open(js_path, "w") as f:
        f.write(js)

    result = subprocess.run(
        ["node", js_path, output_path],
        capture_output=True, text=True, timeout=30
    )
    os.unlink(js_path)

    if result.returncode != 0:
        raise RuntimeError(f"Node.js error: {result.stderr}")

    print(f"  Saved: {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
# HTML report builder
# ─────────────────────────────────────────────────────────────────────────────

def build_html_report(results: list[dict], output_path: str,
                       rfp_file: str, draft_file: str):
    """Generate a standalone interactive HTML compliance report."""

    today     = datetime.date.today().strftime("%B %d, %Y")
    total     = len(results)
    passed    = sum(1 for r in results if r.get("status") == "Pass")
    partial   = sum(1 for r in results if r.get("status") == "Partial")
    missing   = sum(1 for r in results if r.get("status") == "Missing")
    mandatory_missing = sum(1 for r in results
                            if r.get("status") == "Missing" and r.get("mandatory", True))
    score     = round(passed / total * 100) if total else 0

    score_color = "#065F46" if score >= 80 else ("#92400E" if score >= 60 else "#C0392B")

    def status_badge(status):
        colors = {
            "Pass":    ("D1FAE5","065F46"),
            "Partial": ("FEF3C7","92400E"),
            "Missing": ("FDECEA","C0392B"),
            "Unknown": ("F5F5F5","444441"),
        }
        bg, txt = colors.get(status, ("F5F5F5","444441"))
        return (f'<span style="background:#{bg};color:#{txt};padding:3px 10px;'
                f'border-radius:99px;font-size:11px;font-weight:600;'
                f'letter-spacing:.04em">{status.upper()}</span>')

    def req_row(req, i):
        alt = "#F9FAFB" if i % 2 == 0 else "#FFFFFF"
        badge = status_badge(req.get("status","Unknown"))
        mand = ('<span style="color:#C0392B;font-size:11px;font-style:italic">Required</span>'
                if req.get("mandatory", True) else
                '<span style="color:#888;font-size:11px;font-style:italic">Preferred</span>')
        excerpt = req.get("draft_excerpt","")
        exc_html = (f'<div style="margin-top:6px;font-size:11px;color:#666;font-style:italic">'
                    f'Found: {excerpt}</div>'
                    if excerpt and excerpt != "Not found." else "")
        reco = req.get("recommendation","")
        reco_html = (f'<div style="color:#92400E;font-size:12px">{reco}</div>'
                     if reco and reco != "None needed." else
                     '<div style="color:#888;font-size:11px;font-style:italic">None needed</div>')
        source = req.get("source_quote","")
        src_html = (f'<div style="margin-top:4px;font-size:11px;color:#888;font-style:italic">'
                    f'&ldquo;{source}&rdquo;</div>' if source else "")

        return f"""<tr style="background:{alt}">
          <td style="padding:10px 8px;text-align:center;font-weight:600;color:#1A6E6E;
                     border:0.5px solid #E5E7EB;font-size:13px">{req.get('id','')}</td>
          <td style="padding:10px 10px;border:0.5px solid #E5E7EB">
            <div style="font-weight:600;color:#0F6E56;font-size:12px">{req.get('category','')}</div>
            {mand}
          </td>
          <td style="padding:10px 10px;border:0.5px solid #E5E7EB">
            <div style="font-size:13px;color:#1F2937">{req.get('requirement','')}</div>
            {src_html}
          </td>
          <td style="padding:10px 8px;text-align:center;border:0.5px solid #E5E7EB">{badge}</td>
          <td style="padding:10px 10px;border:0.5px solid #E5E7EB">
            <div style="font-size:12px;color:#1F2937">{req.get('finding','')}</div>
            {exc_html}
          </td>
          <td style="padding:10px 10px;border:0.5px solid #E5E7EB">{reco_html}</td>
        </tr>"""

    rows_html = "\n".join(req_row(r, i) for i, r in enumerate(results))

    alert_html = ""
    if mandatory_missing > 0:
        alert_html = f"""<div style="background:#FDECEA;border-left:4px solid #C0392B;
            border-radius:0 8px 8px 0;padding:12px 16px;margin:16px 0;
            font-size:13px;color:#C0392B;font-weight:600">
            ⚠ {mandatory_missing} mandatory requirement(s) are Missing — 
            address these before submitting.
        </div>"""
    else:
        alert_html = """<div style="background:#D1FAE5;border-left:4px solid #1A6E6E;
            border-radius:0 8px 8px 0;padding:12px 16px;margin:16px 0;
            font-size:13px;color:#065F46;font-weight:600">
            ✓ All mandatory requirements are addressed. Review Partial items before submitting.
        </div>"""

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Compliance Report — {ORG_NAME}</title>
<style>
  *{{box-sizing:border-box;margin:0;padding:0}}
  body{{font-family:'Segoe UI',Arial,sans-serif;background:#F8FAFB;color:#1F2937;font-size:14px}}
  .page{{max-width:1300px;margin:0 auto;padding:24px 20px}}
  h1{{font-size:26px;font-weight:700;color:#1A6E6E;margin-bottom:4px}}
  .sub{{font-size:13px;color:#6B7280;margin-bottom:20px}}
  .kpi-row{{display:grid;grid-template-columns:repeat(5,1fr);gap:10px;margin-bottom:16px}}
  .kpi{{background:white;border-radius:10px;padding:14px 16px;text-align:center;
        box-shadow:0 1px 3px rgba(0,0,0,0.06)}}
  .kpi-val{{font-size:28px;font-weight:700;line-height:1}}
  .kpi-lbl{{font-size:11px;color:#6B7280;margin-top:4px}}
  .filters{{display:flex;gap:10px;margin-bottom:14px;flex-wrap:wrap;align-items:center}}
  .filters label{{font-size:12px;color:#6B7280;font-weight:600;text-transform:uppercase;letter-spacing:.06em}}
  select{{padding:6px 12px;border:1px solid #D1D5DB;border-radius:6px;font-size:13px;
          background:white;cursor:pointer}}
  table{{width:100%;border-collapse:collapse;background:white;border-radius:10px;overflow:hidden;
         box-shadow:0 1px 3px rgba(0,0,0,0.06)}}
  th{{background:#1A6E6E;color:white;padding:11px 10px;font-size:12px;font-weight:600;
      text-align:left;border:0.5px solid #0F6E56}}
  .footer{{margin-top:20px;text-align:center;font-size:11px;color:#9CA3AF;padding-top:12px;
           border-top:1px solid #E5E7EB}}
  .hidden{{display:none}}
</style>
</head>
<body>
<div class="page">
  <h1>Grant Compliance Report</h1>
  <div class="sub">{ORG_NAME} &middot; {FUNDER_NAME} &middot; 
    RFP: {rfp_file} &middot; Draft: {draft_file} &middot; {today}</div>

  <div class="kpi-row">
    <div class="kpi"><div class="kpi-val" style="color:#1A6E6E">{total}</div>
      <div class="kpi-lbl">Total requirements</div></div>
    <div class="kpi"><div class="kpi-val" style="color:#065F46">{passed}</div>
      <div class="kpi-lbl">Passing</div></div>
    <div class="kpi"><div class="kpi-val" style="color:#92400E">{partial}</div>
      <div class="kpi-lbl">Partial</div></div>
    <div class="kpi"><div class="kpi-val" style="color:#C0392B">{missing}</div>
      <div class="kpi-lbl">Missing</div></div>
    <div class="kpi"><div class="kpi-val" style="color:{score_color}">{score}%</div>
      <div class="kpi-lbl">Compliance score</div></div>
  </div>

  {alert_html}

  <div class="filters">
    <label>Filter</label>
    <select onchange="filterRows(this.value,'status')">
      <option value="all">All statuses</option>
      <option value="Pass">Pass only</option>
      <option value="Partial">Partial only</option>
      <option value="Missing">Missing only</option>
    </select>
    <select onchange="filterRows(this.value,'category')">
      <option value="all">All categories</option>
      <option>Eligibility</option>
      <option>Program Design</option>
      <option>Evaluation</option>
      <option>Budget</option>
      <option>Organizational Capacity</option>
      <option>Formatting</option>
      <option>Attachments</option>
      <option>Other</option>
    </select>
    <select onchange="filterRows(this.value,'mandatory')">
      <option value="all">Required + Preferred</option>
      <option value="required">Required only</option>
      <option value="preferred">Preferred only</option>
    </select>
  </div>

  <table id="results-table">
    <thead><tr>
      <th style="width:44px">#</th>
      <th style="width:120px">Category</th>
      <th>Requirement</th>
      <th style="width:90px;text-align:center">Status</th>
      <th style="width:240px">What was found</th>
      <th style="width:240px">Recommendation</th>
    </tr></thead>
    <tbody id="results-body">
      {rows_html}
    </tbody>
  </table>

  <div class="footer">
    {ORG_NAME} &middot; Compliance Report &middot; {today} &middot;
    For internal use only — AI-assisted analysis, review before submission
  </div>
</div>

<script>
const rows = document.querySelectorAll('#results-body tr');
let activeStatus = 'all', activeCategory = 'all', activeMandatory = 'all';

function filterRows(val, type) {{
  if(type==='status') activeStatus=val;
  if(type==='category') activeCategory=val;
  if(type==='mandatory') activeMandatory=val;
  rows.forEach(row => {{
    const cells = row.querySelectorAll('td');
    const status = cells[3]?.textContent.trim() || '';
    const category = cells[1]?.textContent.trim().split('\\n')[0].trim() || '';
    const mandatory = cells[1]?.textContent.includes('Required') ? 'required' : 'preferred';
    const show = (activeStatus==='all' || status.includes(activeStatus.toUpperCase())) &&
                 (activeCategory==='all' || category===activeCategory) &&
                 (activeMandatory==='all' || mandatory===activeMandatory);
    row.classList.toggle('hidden', !show);
  }});
}}
</script>
</body>
</html>"""

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"  Saved: {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
# Console summary
# ─────────────────────────────────────────────────────────────────────────────

def print_summary(results: list[dict]):
    total   = len(results)
    passed  = sum(1 for r in results if r.get("status") == "Pass")
    partial = sum(1 for r in results if r.get("status") == "Partial")
    missing = sum(1 for r in results if r.get("status") == "Missing")
    score   = round(passed / total * 100) if total else 0
    mand_missing = [r for r in results if r.get("status")=="Missing" and r.get("mandatory",True)]

    print(f"\n{'═'*62}")
    print(f"  COMPLIANCE REPORT — {ORG_NAME}")
    print(f"{'═'*62}")
    print(f"  Total requirements checked:  {total}")
    print(f"  ✓ Pass:                      {passed}  ({round(passed/total*100)}%)")
    print(f"  ~ Partial:                   {partial}  ({round(partial/total*100)}%)")
    print(f"  ✗ Missing:                   {missing}  ({round(missing/total*100)}%)")
    print(f"  Compliance score:            {score}%")

    if mand_missing:
        print(f"\n  ⚠  MANDATORY ITEMS MISSING ({len(mand_missing)}):")
        for r in mand_missing:
            print(f"     #{r['id']} [{r['category']}] {r['requirement'][:65]}...")
    else:
        print(f"\n  ✓  All mandatory requirements addressed.")

    partial_items = [r for r in results if r.get("status") == "Partial"]
    if partial_items:
        print(f"\n  ~ PARTIAL — needs strengthening ({len(partial_items)}):")
        for r in partial_items[:5]:
            print(f"     #{r['id']} {r['requirement'][:65]}...")
            print(f"          → {r.get('recommendation','')[:80]}")


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

def main():
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise EnvironmentError(
            "ANTHROPIC_API_KEY not set.\n"
            "Run:  export ANTHROPIC_API_KEY='sk-ant-...'"
        )

    print(f"\n📋  Grant Compliance Checker — {ORG_NAME}")
    print(f"    RFP:   {RFP_FILE}")
    print(f"    Draft: {DRAFT_FILE}\n")

    print("  Reading files...")
    rfp_text   = extract_text(RFP_FILE)
    draft_text = extract_text(DRAFT_FILE)
    print(f"  RFP:   {len(rfp_text):,} chars  |  Draft: {len(draft_text):,} chars")

    results = run_compliance_check(rfp_text, draft_text, api_key)

    print_summary(results)

    print(f"\n  Building reports...")
    try:
        build_word_report(results, REPORT_DOCX, RFP_FILE, DRAFT_FILE)
    except Exception as e:
        print(f"  Word report skipped: {e}")

    build_html_report(results, REPORT_HTML, RFP_FILE, DRAFT_FILE)

    print(f"\n✅  Done.")
    print(f"    Open {REPORT_HTML} in any browser for the interactive report.")
    print(f"    Open {REPORT_DOCX} in Word for the formatted document.")


if __name__ == "__main__":
    main()


# ─────────────────────────────────────────────────────────────────────────────
# COLAB INSTRUCTIONS
# ─────────────────────────────────────────────────────────────────────────────
#
# CELL 1 — Install
# !pip install anthropic python-docx PyPDF2 --quiet
#
# CELL 2 — Set API key
# import os
# os.environ["ANTHROPIC_API_KEY"] = "sk-ant-your-key-here"
#
# CELL 3 — Upload your files
# from google.colab import files
# print("Upload your RFP (PDF or DOCX):")
# uploaded = files.upload()
# RFP_FILE = list(uploaded.keys())[0]
#
# print("Upload your draft narrative (DOCX or TXT):")
# uploaded2 = files.upload()
# DRAFT_FILE = list(uploaded2.keys())[0]
#
# CELL 4 — Run the checker
# exec(open('compliance_checker.py').read())
# main()
#
# CELL 5 — Download reports
# from google.colab import files
# files.download('compliance_report.html')
# files.download('compliance_report.docx')   # if Node.js available
